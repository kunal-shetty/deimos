import os
import re
from .base import BaseTool

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

EMPHASIS_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


class CreateDocxTool(BaseTool):
    name = "create_docx"
    description = (
        "Create a Word (.docx) document from a structured content array. "
        "Read the 'docx' skill first (via read_skill) to learn the content "
        "block format and formatting conventions before calling this."
    )
    input_schema = {
        "type": "object",
        "properties": {
            "output_path": {
                "type": "string",
                "description": "File path to save the document to, ending in .docx",
            },
            "content": {
                "type": "array",
                "description": (
                    "Array of content blocks. Each block has a 'type' field: "
                    "'heading' (text, level 1-4), 'paragraph' (text, supports "
                    "**bold**/*italic*), 'bullet_list' (items), "
                    "'numbered_list' (items), 'table' (headers, rows), "
                    "'page_break', or 'spacer'."
                ),
                "items": {"type": "object"},
            },
        },
        "required": ["output_path", "content"],
    }

    def run(self, output_path: str, content: list) -> str:
        if not DOCX_AVAILABLE:
            return "Error: python-docx is not installed. Run: pip install python-docx"

        if not output_path.endswith(".docx"):
            output_path += ".docx"

        try:
            doc = Document()
            for block in content:
                self._render_block(doc, block)

            os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
            doc.save(output_path)
            return f"Created: {os.path.abspath(output_path)} ({len(content)} content blocks)"
        except Exception as e:
            return f"Error creating document: {e}"

    def _render_block(self, doc, block: dict):
        t = block.get("type")

        if t == "heading":
            doc.add_heading(block.get("text", ""), level=max(1, min(4, int(block.get("level", 2)))))

        elif t == "paragraph":
            p = doc.add_paragraph()
            _add_runs(p, block.get("text", ""))

        elif t == "bullet_list":
            for item in block.get("items", []):
                p = doc.add_paragraph(style="List Bullet")
                _add_runs(p, item)

        elif t == "numbered_list":
            for item in block.get("items", []):
                p = doc.add_paragraph(style="List Number")
                _add_runs(p, item)

        elif t == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if not headers:
                return
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            header_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                header_cells[i].text = str(h)
                for p in header_cells[i].paragraphs:
                    for run in p.runs:
                        run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    if i < len(cells):
                        cells[i].text = str(val)

        elif t == "page_break":
            doc.add_page_break()

        elif t == "spacer":
            doc.add_paragraph()


def _add_runs(paragraph, text: str):
    if not text:
        return
    for part in EMPHASIS_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)